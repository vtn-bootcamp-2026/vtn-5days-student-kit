#!/usr/bin/env python3
"""
Capstone Grader Validator — Xác minh tính hợp lệ của kết quả chấm điểm.
Hỗ trợ kiểm định từng bước trong workflow và tự động xuất logs / human review.
"""
from __future__ import annotations

import argparse
import json
import re
import sys
import unicodedata
from pathlib import Path


def log_execution(group_dir: Path, step: int, action: str, status: str, message: str) -> None:
    """Ghi nhật ký thực thi vào cả execution_log.md và executon_log.md dưới dạng bảng Markdown."""
    import datetime
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    log_files = [
        group_dir / "output" / "execution_log.md",
        group_dir / "output" / "executon_log.md"
    ]
    for log_path in log_files:
        try:
            log_path.parent.mkdir(parents=True, exist_ok=True)
            exists = log_path.exists()
            with open(log_path, "a", encoding="utf-8") as f:
                if not exists:
                    f.write(f"---\nmuc-dich: Nhật ký thực thi workflow chấm bài Capstone\ntrang-thai: active\n---\n\n")
                    f.write(f"# Nhật ký thực thi (Execution Log) - {group_dir.name}\n\n")
                    f.write(f"| Thời gian | Bước | Hành động | Trạng thái | Chi tiết |\n")
                    f.write(f"| --- | --- | --- | --- | --- |\n")
                f.write(f"| {timestamp} | Bước {step} | {action} | {status} | {message} |\n")
        except Exception as e:
            print(f"[⚠️] Không thể ghi log thực thi vào {log_path.name}: {e}")


def log_human_review(group_dir: Path, step: int, criteria: str, file: str, quote: str, reason: str, reinit: bool = False) -> None:
    """Ghi nhật ký đánh giá thủ công vào human_review.md. Reinit=True sẽ khởi tạo lại file."""
    import datetime
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    review_path = group_dir / "output" / "human_review.md"
    try:
        review_path.parent.mkdir(parents=True, exist_ok=True)
        exists = review_path.exists()
        
        # Nếu reinit=True, ghi đè file mới. Nếu không, chỉ ghi đè khi chưa tồn tại.
        mode = "w" if (reinit or not exists) else "a"
        
        with open(review_path, mode, encoding="utf-8") as f:
            if mode == "w" or not exists:
                f.write(f"---\nmuc-dich: Danh sách các tiêu chí cần giám khảo con người đánh giá lại\ntrang-thai: active\n---\n\n")
                f.write(f"# Danh sách đánh giá thủ công (Human Review Log) - {group_dir.name}\n\n")
                f.write(f"Vui lòng kiểm tra các mục dưới đây do hệ thống tự động không thể xác minh chính xác hoặc phát hiện sự không tuân thủ.\n\n")
                f.write(f"| Thời gian | Bước | Tiêu chí | File | Minh chứng (Quote) | Lý do cần review |\n")
                f.write(f"| --- | --- | --- | --- | --- | --- |\n")
            
            quote_esc = quote.replace("|", "\\|") if quote else "N/A"
            reason_esc = reason.replace("|", "\\|") if reason else ""
            f.write(f"| {timestamp} | Bước {step} | {criteria} | `{file}` | `\"{quote_esc}\"` | {reason_esc} |\n")
    except Exception as e:
        print(f"[⚠️] Không thể ghi nhận human review: {e}")


def find_workspace_file(group_dir: Path, rel_path: str) -> Path | None:
    """Tìm tệp tin trong toàn bộ phân cấp thư mục liên quan (nhóm, capstone_project, session root, v.v.)."""
    for base in [group_dir, group_dir / "capstone_project", group_dir.parent, group_dir.parent.parent, group_dir.parent.parent.parent]:
        p = base / rel_path
        if p.exists():
            return p
    return None


def load_json_schema(schema_path: Path) -> dict | None:
    try:
        with open(schema_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        print(f"[⚠️] Lỗi khi load schema '{schema_path.name}': {e}")
        return None


def validate_schema(data: dict, schema_path: Path) -> bool:
    schema = load_json_schema(schema_path)
    if not schema:
        return True  # Bỏ qua nếu không load được schema
        
    try:
        import jsonschema
        jsonschema.validate(instance=data, schema=schema)
        print(f"[OK] Dữ liệu JSON hợp lệ theo schema '{schema_path.name}'.")
        return True
    except ImportError:
        print("[ℹ️] Chưa cài đặt thư viện 'jsonschema'. Bỏ qua kiểm tra cấu trúc JSON schema.")
        return True
    except Exception as e:
        print(f"[FAIL] Dữ liệu JSON không hợp lệ theo schema '{schema_path.name}': {e}")
        return False


def validate_step_1(group_path: Path) -> bool:
    """Kiểm định Bước 1: Tiếp nhận và Thu thập (submission_status.json)"""
    log_execution(group_path, 1, "Bắt đầu kiểm định Bước 1 (Intake)", "INFO", "Xác minh file submission_status.json.")
    
    status_file = group_path / "output" / "submission_status.json"
    if not status_file.exists():
        msg = "Tệp hiện trạng bài nộp output/submission_status.json không tồn tại."
        log_execution(group_path, 1, "Kiểm định Bước 1", "FAIL", msg)
        print(f"[FAIL] {msg}")
        return False

    # Khởi tạo lại human_review.md tại bước kiểm định đầu tiên để làm sạch log cũ
    log_human_review(group_path, 1, "Hệ thống", "N/A", "N/A", "Khởi tạo tiến trình kiểm định mới.", reinit=True)

    try:
        status_data = json.loads(status_file.read_text(encoding="utf-8"))
    except Exception as e:
        msg = f"Không thể giải mã tệp JSON hiện trạng: {e}"
        log_execution(group_path, 1, "Kiểm định Bước 1", "FAIL", msg)
        print(f"[FAIL] {msg}")
        return False

    script_dir = Path(__file__).resolve().parent
    schema_path = script_dir.parent / "schemas" / "capstone-group.schema.json"
    schema_ok = validate_schema(status_data, schema_path)
    if not schema_ok:
        log_execution(group_path, 1, "Kiểm định Bước 1", "FAIL", "Cấu trúc JSON không khớp schema.")
        return False

    # Kiểm tra tính đầy đủ của các tệp và placeholders
    all_ok = True
    # 1. Blueprint files
    for key, info in status_data.get("blueprint_files", {}).items():
        # map exact name for better display
        mappings = {
            "one_pager": "01-use-case-one-pager.md",
            "logical_workflow": "02-logical-workflow.md",
            "core_prompt_design": "03-core-prompt-design.md",
            "compliance_checklist": "04-compliance-checklist.md",
            "action_plan": "05-action-plan-30-90-days.md"
        }
        display_name = info.get("resolved_path") or mappings.get(key, f"{key}.md")
        
        if not info.get("exists", False):
            reason = f"Tài liệu thiết kế quan trọng chưa được nộp."
            log_human_review(group_path, 1, f"Blueprint: {key}", display_name, "N/A", reason)
            all_ok = False
        elif info.get("contains_placeholders", False):
            reason = f"Tài liệu thiết kế còn chứa placeholder chưa hoàn thiện."
            log_human_review(group_path, 1, f"Blueprint: {key}", display_name, "N/A", reason)
            all_ok = False
        
        # Cổng chất lượng: kiểm tra confidence score của từng file
        conf = info.get("confidence_score", 1.0)
        if conf < 0.8:
            reason = f"Độ tin cậy thu thập tệp thấp ({conf:.2f}) do chứa placeholder hoặc chưa cấu hình chuẩn."
            log_human_review(group_path, 1, f"Blueprint: {key} (Confidence)", display_name, "N/A", reason)
            all_ok = False

    # 2. Source code
    anonymizer_info = status_data.get("implementation_files", {}).get("anonymizer_py", {})
    source_summary = status_data.get("implementation_files", {}).get("source_summary", {})
    code_file_count = int(source_summary.get("code_file_count", 0))
    if not anonymizer_info.get("exists", False) and code_file_count == 0:
        reason = "Không phát hiện mã nguồn triển khai trong bài nộp."
        log_human_review(group_path, 1, "Mã nguồn", "N/A", "N/A", reason)
        all_ok = False

    # Kiểm tra độ tin cậy tổng thể của bước Intake
    intake_conf = status_data.get("confidence_score", 1.0)
    if intake_conf < 0.85:
        msg = f"Độ tin cậy tổng thể của hồ sơ bài nộp thấp ({intake_conf:.2f} < 0.85)."
        log_execution(group_path, 1, "Cổng chất lượng Intake", "WARNING", msg)
        log_human_review(group_path, 1, "Tổng thể Intake", "submission_status.json", "N/A", msg + " Yêu cầu giám khảo rà soát.")

    if all_ok:
        log_execution(group_path, 1, "Kiểm định Bước 1", "SUCCESS", "Tất cả các tài liệu đã nộp đầy đủ và không chứa placeholder.")
        print("[OK] Bước 1 kiểm định thành công!")
        return True
    else:
        log_execution(group_path, 1, "Kiểm định Bước 1", "WARNING", "Phát hiện thiếu tài liệu hoặc chứa placeholder/độ tin cậy thấp. Đã log vào human_review.md.")
        print("[WARNING] Phát hiện thiếu tài liệu hoặc chứa placeholder/độ tin cậy thấp. Chi tiết ghi trong human_review.md.")
        # Vẫn cho qua để chạy bước tiếp theo, nhưng báo warning
        return True


def validate_step_2(group_path: Path) -> bool:
    """Kiểm định Bước 2: Chấm điểm và Trích xuất minh chứng (grading_result_draft.json)"""
    log_execution(group_path, 2, "Bắt đầu kiểm định Bước 2 (Assessment)", "INFO", "Xác minh file grading_result_draft.json.")
    
    grading_file = group_path / "output" / "grading_result_draft.json"
    if not grading_file.exists():
        msg = "Tệp chấm điểm dự thảo output/grading_result_draft.json không tồn tại."
        log_execution(group_path, 2, "Kiểm định Bước 2", "FAIL", msg)
        print(f"[FAIL] {msg}")
        return False

    try:
        grading_data = json.loads(grading_file.read_text(encoding="utf-8"))
    except Exception as e:
        msg = f"Không thể giải mã tệp JSON kết quả chấm điểm: {e}"
        log_execution(group_path, 2, "Kiểm định Bước 2", "FAIL", msg)
        print(f"[FAIL] {msg}")
        return False

    script_dir = Path(__file__).resolve().parent
    schema_path = script_dir.parent / "schemas" / "grading-result.schema.json"
    schema_ok = validate_schema(grading_data, schema_path)
    if not schema_ok:
        log_execution(group_path, 2, "Kiểm định Bước 2", "FAIL", "Cấu trúc JSON không khớp schema.")
        return False

    # Xác minh minh chứng thực tế khớp 100%
    evidence_ok = True
    evidences = []
    
    # Gom minh chứng
    for criteria_key, criteria_val in grading_data.get("main_criteria", {}).items():
        score = criteria_val.get("score", 0)
        # Kiểm tra nếu điểm tiêu chí bằng 0
        if score == 0:
            log_human_review(group_path, 2, f"Main Criteria: {criteria_key}", "N/A", "N/A", "Tiêu chí chính nhận điểm 0, cần giám khảo đánh giá lại.")
            
        for ev in criteria_val.get("evidence", []):
            evidences.append((f"Main Criteria: {criteria_key}", ev))
            
    for bonus_key, bonus_val in grading_data.get("bonus_criteria", {}).items():
        for ev in bonus_val.get("evidence", []):
            evidences.append((f"Bonus Criteria: {bonus_key}", ev))

    for source_key, ev in evidences:
        filename = ev.get("file")
        quote = ev.get("quote")
        
        if not filename or not quote:
            reason = f"Minh chứng thiếu thông tin file hoặc quote."
            log_human_review(group_path, 2, source_key, filename or "N/A", quote or "N/A", reason)
            evidence_ok = False
            continue

        # Định vị file
        file_path = find_workspace_file(group_path, filename)

        if not file_path or not file_path.exists():
            reason = f"Tệp chứa minh chứng không tồn tại thực tế."
            log_human_review(group_path, 2, source_key, filename, quote, reason)
            evidence_ok = False
            continue

        # Đọc nội dung
        try:
            if filename.lower().endswith(('.pdf', '.pptx', '.docx', '.zip')):
                print(f"[OK] Xác minh tệp nhị phân {filename} tồn tại thực tế.")
                continue
            file_content = file_path.read_text(encoding="utf-8")
        except Exception as e:
            reason = f"Không thể đọc tệp để kiểm tra minh chứng: {e}"
            log_human_review(group_path, 2, source_key, filename, quote, reason)
            evidence_ok = False
            continue


        # Chuẩn hóa unicode (NFC)
        file_content_norm = unicodedata.normalize("NFC", file_content)
        quote_norm = unicodedata.normalize("NFC", quote)

        # Chuẩn hóa khoảng trắng
        file_content_clean = re.sub(r"\s+", " ", file_content_norm)
        quote_clean = re.sub(r"\s+", " ", quote_norm)

        # Kiểm tra quote mặc định
        if "Nội dung phản ánh yêu cầu đạt" in quote or quote.strip() == "":
            reason = "Vẫn dùng minh chứng (quote) mặc định chưa được cập nhật."
            log_human_review(group_path, 2, source_key, filename, quote, reason)
            evidence_ok = False
            continue

        if quote_clean not in file_content_clean:
            reason = "Đoạn trích dẫn (quote) minh chứng không khớp với nội dung thực tế của tệp."
            log_human_review(group_path, 2, source_key, filename, quote, reason)
            evidence_ok = False
            print(f"[FAIL] Không khớp minh chứng cho {source_key} trong file {filename}.")
        else:
            print(f"[OK] Khớp minh chứng cho {source_key} trong file {filename}.")

    # 3. Cổng chất lượng (Quality Gate) về Độ tin cậy (Confidence Score)
    quality_gate_ok = True
    score_gate_ok = True
    total_conf = 0.0
    criteria_count = 0
    max_scores = {
        "Main Criteria: blueprint_overall": 30.0,
        "Main Criteria: logical_workflow": 20.0,
        "Main Criteria: core_prompt": 20.0,
        "Main Criteria: action_plan": 15.0,
        "Main Criteria: presentation": 15.0,
        "Bonus Criteria: source_code_testing": 10.0,
        "Bonus Criteria: safety_logging": 10.0,
    }
    
    # Gom danh sách tiêu chí để duyệt
    all_criteria = []
    for k, v in grading_data.get("main_criteria", {}).items():
        all_criteria.append((f"Main Criteria: {k}", v))
    for k, v in grading_data.get("bonus_criteria", {}).items():
        all_criteria.append((f"Bonus Criteria: {k}", v))
        
    for name, val in all_criteria:
        score = float(val.get("score", 0.0))
        max_score = max_scores.get(name)
        if score < 0 or (max_score is not None and score > max_score):
            msg = f"Điểm tiêu chí '{name}' nằm ngoài thang rubric ({score} / {max_score})."
            print(f"[FAIL] {msg}")
            log_human_review(group_path, 2, name, "grading_result_draft.json", str(score), msg)
            log_execution(group_path, 2, "Cổng chất lượng điểm số", "FAIL", msg)
            score_gate_ok = False

        conf_score = val.get("confidence_score", 1.0)
        total_conf += conf_score
        criteria_count += 1
        
        # Ngưỡng từ chối hoàn toàn (Reject - Thấp dưới 0.6)
        if conf_score < 0.6:
            msg = f"Từ chối kết quả chấm (REJECT): Độ tin cậy quá thấp ({conf_score:.2f} < 0.60) ở tiêu chí '{name}'."
            print(f"[FAIL] {msg}")
            log_human_review(group_path, 2, name, "N/A", "N/A", f"Từ chối kết quả chấm (REJECT): Độ tin cậy quá thấp ({conf_score:.2f} < 0.60).")
            log_execution(group_path, 2, "Cổng chất lượng chấm điểm", "FAIL", msg)
            quality_gate_ok = False
        # Ngưỡng cần đánh giá thủ công (Need Review - Dưới 0.85)
        elif conf_score < 0.85:
            msg = f"Yêu cầu đánh giá lại (NEED_REVIEW): Độ tin cậy trung bình ({conf_score:.2f} < 0.85) ở tiêu chí '{name}'."
            print(f"[WARNING] {msg}")
            log_human_review(group_path, 2, name, "N/A", "N/A", f"Yêu cầu đánh giá lại (NEED_REVIEW): Độ tin cậy chưa đạt tối đa ({conf_score:.2f} < 0.85).")
            log_execution(group_path, 2, "Cổng chất lượng chấm điểm", "WARNING", msg)

    if criteria_count > 0:
        avg_conf = total_conf / criteria_count
        if avg_conf < 0.85:
            msg = f"Độ tin cậy chấm điểm trung bình thấp ({avg_conf:.2f} < 0.85)."
            print(f"[WARNING] {msg}")
            log_human_review(group_path, 2, "Tổng thể chấm điểm", "N/A", "N/A", f"Độ tin cậy chấm điểm trung bình thấp ({avg_conf:.2f} < 0.85). Cần giám khảo đánh giá lại.")
            log_execution(group_path, 2, "Cổng chất lượng chấm điểm tổng thể", "WARNING", msg)

    main_total = sum(float(v.get("score", 0.0)) for v in grading_data.get("main_criteria", {}).values())
    bonus_total = sum(float(v.get("score", 0.0)) for v in grading_data.get("bonus_criteria", {}).values())
    reported_main_total = float(grading_data.get("total_score", 0.0))
    reported_bonus_total = float(grading_data.get("bonus_score", bonus_total))
    reported_total_with_bonus = float(grading_data.get("total_score_with_bonus", reported_main_total + reported_bonus_total))
    if abs(main_total - reported_main_total) > 0.05:
        msg = f"Tổng điểm chính không khớp chi tiết ({reported_main_total} != {main_total:.1f})."
        print(f"[FAIL] {msg}")
        log_human_review(group_path, 2, "Tổng điểm chính", "grading_result_draft.json", str(reported_main_total), msg)
        log_execution(group_path, 2, "Cổng chất lượng điểm số", "FAIL", msg)
        score_gate_ok = False
    if abs(bonus_total - reported_bonus_total) > 0.05:
        msg = f"Tổng điểm cộng không khớp chi tiết ({reported_bonus_total} != {bonus_total:.1f})."
        print(f"[FAIL] {msg}")
        log_human_review(group_path, 2, "Tổng điểm cộng", "grading_result_draft.json", str(reported_bonus_total), msg)
        log_execution(group_path, 2, "Cổng chất lượng điểm số", "FAIL", msg)
        score_gate_ok = False
    if abs((main_total + bonus_total) - reported_total_with_bonus) > 0.05:
        msg = f"Tổng điểm có điểm cộng không khớp chi tiết ({reported_total_with_bonus} != {main_total + bonus_total:.1f})."
        print(f"[FAIL] {msg}")
        log_human_review(group_path, 2, "Tổng điểm có điểm cộng", "grading_result_draft.json", str(reported_total_with_bonus), msg)
        log_execution(group_path, 2, "Cổng chất lượng điểm số", "FAIL", msg)
        score_gate_ok = False

    if evidence_ok and quality_gate_ok and score_gate_ok:
        log_execution(group_path, 2, "Kiểm định Bước 2", "SUCCESS", "Kết quả chấm điểm khớp 100% minh chứng thực tế và vượt qua cổng chất lượng độ tin cậy.")
        print("[OK] Bước 2 kiểm định thành công!")
        return True
    else:
        status_msg = "Phát hiện lỗi kiểm định: "
        if not evidence_ok:
            status_msg += "sai lệch minh chứng trích dẫn; "
        if not quality_gate_ok:
            status_msg += "độ tin cậy chấm điểm dưới ngưỡng tối thiểu (REJECT);"
        if not score_gate_ok:
            status_msg += "điểm số không khớp thang rubric hoặc tổng điểm;"
        log_execution(group_path, 2, "Kiểm định Bước 2", "FAIL", status_msg)
        print(f"[FAIL] {status_msg} Chi tiết ghi trong human_review.md.")
        return False


def validate_step_4(group_path: Path) -> bool:
    """Kiểm định Bước 4: Xuất báo cáo DOCX (kiểm định tệp docx kết quả)"""
    log_execution(group_path, 4, "Bắt đầu kiểm định Bước 4 (DOCX)", "INFO", "Xác minh file báo cáo Word kết quả.")
    
    docx_path = group_path / "output" / f"danh_gia_{group_path.name}.docx"
    if not docx_path.exists():
        msg = f"Tệp báo cáo Word output/danh_gia_{group_path.name}.docx không tồn tại."
        log_execution(group_path, 4, "Kiểm định Bước 4", "FAIL", msg)
        log_human_review(group_path, 4, "Báo cáo Word", f"danh_gia_{group_path.name}.docx", "N/A", msg)
        print(f"[FAIL] {msg}")
        return False

    # Đọc cấu trúc file Word để kiểm định
    try:
        from docx import Document
        doc = Document(str(docx_path))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        tables_count = len(doc.tables)
        
        if len(headings) >= 3 and tables_count >= 2:
            # Kiểm tra số cột của các bảng
            for table in doc.tables:
                if len(table.columns) < 5:
                    reason = f"Bảng trong file Word không đủ 5 cột (có {len(table.columns)} cột)."
                    log_execution(group_path, 4, "Kiểm định Bước 4", "FAIL", reason)
                    log_human_review(group_path, 4, "Báo cáo Word", f"danh_gia_{group_path.name}.docx", "N/A", reason)
                    print(f"[FAIL] {reason}")
                    return False
            log_execution(group_path, 4, "Kiểm định Bước 4", "SUCCESS", "Báo cáo Word hợp lệ về mặt cấu trúc (đủ bảng biểu 5 cột và tiêu đề).")
            print("[OK] Bước 4 kiểm định thành công!")
            return True
        else:
            reason = f"File Word thiếu các thành phần yêu cầu (có {len(headings)} headings, {tables_count} tables)."
            log_execution(group_path, 4, "Kiểm định Bước 4", "FAIL", reason)
            log_human_review(group_path, 4, "Báo cáo Word", f"danh_gia_{group_path.name}.docx", "N/A", reason)
            print(f"[FAIL] {reason}")
            return False
    except ImportError:
        # Nếu chưa cài python-docx ở môi trường chạy validator, log thông báo cảnh báo nhẹ
        msg = "Chưa cài đặt python-docx trong môi trường chạy validator. Tạm thời bỏ qua đọc cấu trúc chi tiết."
        log_execution(group_path, 4, "Kiểm định Bước 4", "WARNING", msg)
        print(f"[WARNING] {msg}")
        return True
    except Exception as e:
        reason = f"Lỗi đọc hoặc phân tích file Word: {e}"
        log_execution(group_path, 4, "Kiểm định Bước 4", "FAIL", reason)
        log_human_review(group_path, 4, "Báo cáo Word", f"danh_gia_{group_path.name}.docx", "N/A", reason)
        print(f"[FAIL] {reason}")
        return False


def main() -> None:
    parser = argparse.ArgumentParser(description="Capstone Grader Validator")
    parser.add_argument("group_path", type=str, help="Đường dẫn đến thư mục của nhóm")
    parser.add_argument("-s", "--step", type=int, choices=[1, 2, 4], default=None, 
                        help="Chỉ định bước cụ thể để kiểm định (1: Intake, 2: Assessment, 4: DOCX). Nếu để trống, kiểm tra tất cả các bước.")
    args = parser.parse_args()

    group_path = Path(args.group_path).resolve()
    if not group_path.exists() or not group_path.is_dir():
        print(f"Lỗi: Thư mục nhóm '{group_path}' không tồn tại.")
        sys.exit(1)

    print(f"=== KHỞI CHẠY KIỂM ĐỊNH CHO NHÓM: {group_path.name} ===")
    
    success = True
    
    if args.step is None:
        # Tự động chạy tất cả các bước có sẵn tệp tương ứng
        status_file = group_path / "output" / "submission_status.json"
        grading_file = group_path / "output" / "grading_result_draft.json"
        docx_path = group_path / "output" / f"danh_gia_{group_path.name}.docx"
        
        run_any = False
        if status_file.exists():
            run_any = True
            success = success and validate_step_1(group_path)
        if grading_file.exists():
            run_any = True
            success = success and validate_step_2(group_path)
        if docx_path.exists():
            run_any = True
            success = success and validate_step_4(group_path)
            
        if not run_any:
            print("[⚠️] Không tìm thấy bất kỳ tệp kết quả nào để kiểm định trong thư mục output.")
            sys.exit(1)
    else:
        if args.step == 1:
            success = validate_step_1(group_path)
        elif args.step == 2:
            success = validate_step_2(group_path)
        elif args.step == 4:
            success = validate_step_4(group_path)

    print("---------------------------------------------")
    if success:
        print("[OK] Tiến trình kiểm định hoàn tất thành công!")
        sys.exit(0)
    else:
        print("[FAIL] Phát hiện lỗi kiểm định. Vui lòng kiểm tra lại log và tệp human_review.md.")
        sys.exit(1)


if __name__ == "__main__":
    main()
