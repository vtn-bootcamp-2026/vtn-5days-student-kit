#!/usr/bin/env python3
"""
Capstone Report Generator — Xuất báo cáo chấm điểm Capstone dạng file Word (DOCX).
Tự động cài đặt python-docx nếu chưa tồn tại và tạo tệp DOCX đẹp mắt, chuyên nghiệp.
"""
from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path

# Thử import python-docx, nếu thiếu sẽ tự cài đặt
try:
    import docx
except ImportError:
    print("[ℹ️] Không tìm thấy thư viện 'python-docx'. Đang tự động cài đặt...")
    try:
        subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], check=True)
        import docx
        print("[OK] Cài đặt 'python-docx' thành công!")
    except Exception as e:
        print(f"[FAIL] Không thể tự động cài đặt 'python-docx': {e}")
        print("Vui lòng cài đặt thủ công bằng lệnh: pip install python-docx")
        sys.exit(1)

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def build_docx_report(grading_data: dict, output_path: Path) -> None:
    """Tạo tệp báo cáo DOCX từ dữ liệu chấm điểm JSON."""
    doc = Document()
    
    # 1. Cấu hình định dạng trang và Font chữ
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # 2. Tiêu đề báo cáo
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("BÁO CÁO ĐÁNH GIÁ VÀ CHẤM ĐIỂM DỰ ÁN CAPSTONE")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(180, 0, 0) # Màu đỏ đậm thương hiệu

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(f"Nhóm đánh giá: {grading_data.get('group_id', 'N/A').upper()} — Khóa học AI Builders Bootcamp")
    sub_run.italic = True
    sub_run.font.size = Pt(12)

    doc.add_paragraph() # Dòng trống

    # 3. Thông tin chung
    doc.add_heading("I. Thông tin chung", level=2)
    p_info = doc.add_paragraph()
    p_info.add_run(f"• Giám khảo chấm: ").bold = True
    p_info.add_run(f"{grading_data.get('grader_name', 'N/A')}\n")
    p_info.add_run(f"• Ngày đánh giá: ").bold = True
    p_info.add_run(f"{grading_data.get('evaluation_date', 'N/A')}\n")
    p_info.add_run(f"• Tổng điểm phần chấm chính: ").bold = True
    p_info.add_run(f"{grading_data.get('total_score', 0)} / 100 điểm\n")
    p_info.add_run(f"• Xếp loại kết quả: ").bold = True
    p_info.add_run(f"{grading_data.get('final_classification', 'N/A')}").bold = True

    doc.add_paragraph()

    # 4. Bảng điểm chi tiết phần chấm chính
    doc.add_heading("II. Điểm số chi tiết phần chấm chính (Thang 100 điểm)", level=2)
    
    # Tạo bảng
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Shading Accent 1'
    
    # Thiết lập header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tiêu chí đánh giá'
    hdr_cells[1].text = 'Điểm đạt'
    hdr_cells[2].text = 'Độ tin cậy'
    hdr_cells[3].text = 'Minh chứng trích dẫn (Evidence)'
    hdr_cells[4].text = 'Nhận xét chi tiết'

    # Map key sang Tiếng Việt
    main_mappings = {
        "blueprint_overall": "1. Biểu mẫu thiết kế giải pháp: blueprint",
        "logical_workflow": "2. Luồng logic hợp lý",
        "core_prompt": "3. Lời nhắc bền vững",
        "action_plan": "4. Lộ trình khả thi",
        "presentation": "5. Slide & thuyết trình chuyên nghiệp"
    }

    main_crit = grading_data.get("main_criteria", {})
    for key, label in main_mappings.items():
        if key in main_crit:
            val = main_crit[key]
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = str(val.get("score", 0))
            
            # Ghi nhận độ tin cậy
            conf = val.get("confidence_score", 1.0)
            conf_text = f"{conf*100:.1f}%"
            if conf < 0.6:
                conf_text += " (REJECT)"
            elif conf < 0.85:
                conf_text += " (Review)"
            row_cells[2].text = conf_text
            
            # Gộp các quotes
            quotes = []
            for ev in val.get("evidence", []):
                quotes.append(f"• [{ev.get('file')}]: \"{ev.get('quote')}\"")
            row_cells[3].text = "\n".join(quotes)
            row_cells[4].text = val.get("comments", "")

    doc.add_paragraph()

    # 5. Bảng điểm cộng khuyến khích
    doc.add_heading("III. Phần điểm cộng khuyến khích (Tối đa +20 điểm)", level=2)
    
    table_bonus = doc.add_table(rows=1, cols=5)
    table_bonus.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_bonus.style = 'Light Shading Accent 1'
    
    hdr_cells_bonus = table_bonus.rows[0].cells
    hdr_cells_bonus[0].text = 'Hạng mục cộng điểm'
    hdr_cells_bonus[1].text = 'Điểm cộng'
    hdr_cells_bonus[2].text = 'Độ tin cậy'
    hdr_cells_bonus[3].text = 'Minh chứng trích dẫn (Evidence)'
    hdr_cells_bonus[4].text = 'Nhận xét chi tiết'

    bonus_mappings = {
        "source_code_testing": "Cộng điểm 1: Triển khai mã nguồn & kiểm thử",
        "safety_logging": "Cộng điểm 2: An toàn bảo mật & logs vận hành"
    }

    bonus_crit = grading_data.get("bonus_criteria", {})
    for key, label in bonus_mappings.items():
        if key in bonus_crit:
            val = bonus_crit[key]
            row_cells = table_bonus.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = f"+{val.get('score', 0)}"
            
            # Ghi nhận độ tin cậy
            conf = val.get("confidence_score", 1.0)
            conf_text = f"{conf*100:.1f}%"
            if conf < 0.6:
                conf_text += " (REJECT)"
            elif conf < 0.85:
                conf_text += " (Review)"
            row_cells[2].text = conf_text
            
            quotes = []
            for ev in val.get("evidence", []):
                quotes.append(f"• [{ev.get('file')}]: \"{ev.get('quote')}\"")
            row_cells[3].text = "\n".join(quotes)
            row_cells[4].text = val.get("comments", "")

    doc.add_paragraph()

    # 6. Nhận xét tổng quát
    doc.add_heading("IV. Nhận xét tổng quát và Kết luận", level=2)
    doc.add_paragraph(grading_data.get("general_comment", ""))

    # 7. Chữ ký giám khảo
    doc.add_paragraph()
    sig_p = doc.add_paragraph()
    sig_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_p.add_run("Giám khảo đánh giá\n").bold = True
    sig_p.add_run("(Ký và ghi rõ họ tên)\n\n\n\n")
    sig_p.add_run(grading_data.get("grader_name", "Giám khảo AI")).italic = True

    # Lưu file
    doc.save(str(output_path))
    print(f"[OK] Đã xuất báo cáo DOCX thành công tại: {output_path}")


def validate_generated_docx(docx_path: Path) -> bool:
    """Xác minh xem tệp DOCX được tạo ra có hợp lệ và đọc được không."""
    try:
        doc = Document(str(docx_path))
        # Kiểm tra xem có dữ liệu không
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        tables_count = len(doc.tables)
        
        if len(headings) >= 3 and tables_count >= 2:
            # Kiểm tra số cột của các bảng
            for table in doc.tables:
                if len(table.columns) < 5:
                    print(f"[FAIL] Bảng trong DOCX không đủ 5 cột (có {len(table.columns)} cột).")
                    return False
            print("[OK] Kiểm định cấu trúc file DOCX: Đạt yêu cầu (có đầy đủ tiêu đề và bảng điểm 5 cột).")
            return True
        else:
            print("[FAIL] File DOCX thiếu một số thành phần quan trọng.")
            return False
    except Exception as e:
        print(f"[FAIL] Lỗi kiểm định file DOCX: {e}")
        return False


def log_execution(group_dir: Path, step: int, action: str, status: str, message: str) -> None:
    """Ghi nhật ký thực thi vào cả execution_log.md và executon_log.md dưới dạng bảng Markdown."""
    import datetime
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    log_files = [
        group_dir / "output" / "execution_log.md",
        group_dir / "output" / "executon_log.md"
    ]
    for log_path in log_files:
        try:
            log_path.parent.mkdir(parents=True, exist_ok=True)
            exists = log_path.exists()
            with open(log_path, "a", encoding="utf-8") as f:
                if not exists:
                    f.write(f"---\nmuc-dich: Nhật ký thực thi workflow chấm bài Capstone\ntrang-thai: active\n---\n\n")
                    f.write(f"# Nhật ký thực thi (Execution Log) - {group_dir.name}\n\n")
                    f.write(f"| Thời gian | Bước | Hành động | Trạng thái | Chi tiết |\n")
                    f.write(f"| --- | --- | --- | --- | --- |\n")
                f.write(f"| {timestamp} | Bước {step} | {action} | {status} | {message} |\n")
        except Exception as e:
            print(f"[⚠️] Không thể ghi log thực thi vào {log_path.name}: {e}")


def main() -> None:
    if len(sys.argv) < 2:
        print("Sử dụng: python docx_generator.py <duong_dan_nhom>")
        sys.exit(1)

    group_path = Path(sys.argv[1]).resolve()
    if not group_path.exists() or not group_path.is_dir():
        print(f"Lỗi: Thư mục nhóm '{group_path}' không tồn tại.")
        sys.exit(1)

    # Đảm bảo thư mục output tồn tại
    output_dir = group_path / "output"
    output_dir.mkdir(parents=True, exist_ok=True)

    log_execution(group_path, 4, "Khởi chạy Docx Generator", "INFO", "Bắt đầu xuất báo cáo đánh giá định dạng Word (DOCX).")

    grading_file = output_dir / "grading_result_draft.json"
    if not grading_file.exists():
        msg = f"Chưa có file chấm điểm dự thảo 'output/{grading_file.name}' trong thư mục nhóm."
        log_execution(group_path, 4, "Khởi chạy Docx Generator", "FAIL", msg)
        print(f"Lỗi: {msg}")
        sys.exit(1)

    with open(grading_file, "r", encoding="utf-8") as f:
        grading_data = json.load(f)
    
    docx_output = output_dir / f"danh_gia_{group_path.name}.docx"

    # Tạo tệp DOCX
    try:
        build_docx_report(grading_data, docx_output)
        log_execution(group_path, 4, "Khởi tạo tài liệu DOCX", "SUCCESS", f"Đã sinh tệp báo cáo thành công tại output/danh_gia_{group_path.name}.docx.")
    except Exception as e:
        log_execution(group_path, 4, "Khởi tạo tài liệu DOCX", "FAIL", f"Lỗi trong quá trình sinh báo cáo Word: {e}")
        print(f"Lỗi: {e}")
        sys.exit(1)
    
    # Kiểm định tệp DOCX
    docx_ok = validate_generated_docx(docx_output)
    
    if docx_ok:
        log_execution(group_path, 4, "Kiểm định cấu trúc DOCX nội bộ", "SUCCESS", "Tệp Word hợp lệ về mặt cấu trúc.")
        sys.exit(0)
    else:
        log_execution(group_path, 4, "Kiểm định cấu trúc DOCX nội bộ", "FAIL", "Tệp Word không hợp lệ hoặc thiếu thành phần.")
        sys.exit(1)

if __name__ == "__main__":
    main()
